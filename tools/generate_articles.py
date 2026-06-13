#!/usr/bin/env python3
"""Generate article and series pages from the approved Word source."""

from __future__ import annotations

import html
import json
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT.parent / "專欄文章" / "專欄文章.docx"
WORKFLOW = ROOT.parent / "WorkFlow" / "project" / "YucheinHomePage" / "docs"
AUTHOR = "黃郁倩 諮商心理師"
SITE = "https://yuchienpsy.com"
LINE_URL = "https://line.me/R/ti/p/@264kulgk"


@dataclass
class Series:
    key: str
    name: str
    topic: str
    description: str
    note: str
    image_stem: str
    image_alt: str
    subtitle: str = ""


SERIES = {
    "addiction": Series(
        "addiction",
        "「癮」未條",
        "成癮與依賴",
        "從短影音、咖啡因、孤獨、拖延、孩子黏螢幕到 AI 依賴，理解停不下來背後的心理需求。",
        "先理解行為如何幫助我們撐過壓力，才有機會慢慢找回選擇。",
        "ill-002-addiction-series",
        "一位成人坐在沙發上看向窗外，手中手機旁有柔和循環線條，象徵停不下來的習慣與重新選擇。",
    ),
    "workplace": Series(
        "workplace",
        "咖啡哪有工作苦",
        "職場壓力與自我照顧",
        "從職家平衡、職場 PUA、拒絕、霸凌、職業倦怠到同情疲勞，看見工作裡說不出口的累。",
        "工作困境不只是效率問題，也與界線、安全感、自我價值及身心負荷有關。",
        "ill-003-workplace-series",
        "一位上班族坐在工作桌前伸展肩膀，桌上有筆電、咖啡與待辦紙張，窗邊植物帶來休息感。",
    ),
    "boundary": Series(
        "boundary",
        "「性騷擾」是什麼？",
        "界線與心理安全",
        "從身體自主權、跟騷、數位性別暴力、職場與權勢性騷擾，重新理解關係中的界線。",
        "不舒服的感受值得被理解；認識權力與界線，是保護自己與他人的開始。",
        "ill-004-boundary-series",
        "一位成人以溫和但清楚的手勢表達界線，周圍有柔和光線與支持者的身影。",
    ),
    "self-care": Series(
        "self-care",
        "自我理解與心理照顧",
        "自我理解與照顧",
        "從心理觀點談脆弱、勇敢、自我接納與真實關係。",
        "脆弱不是失敗，而是我們更誠實地理解自己、靠近關係的入口。",
        "ill-005-self-care-series",
        "一位成人坐在扶手椅上輕抱自己，胸前有柔和光點，象徵接納脆弱與自我照顧。",
        "心理師的心靈雞湯",
    ),
    "insights": Series(
        "insights",
        "其他文章／心理觀點",
        "生活與關係心理",
        "收錄關係、自我照顧與社會事件相關的心理觀點文章。",
        "生活中的關係與事件會牽動情緒；理解反應背後的需要，能幫助我們重新安頓自己。",
        "ill-006-insights-series",
        "三個生活片段以紙頁拼貼呈現，包含關係對話、社會事件後的休息與獨處書寫。",
    ),
}

DISCLAIMER = "本文為心理健康科普資訊，不能取代個別診斷、治療或緊急協助。"
ECARE_URL = "https://ecare.mohw.gov.tw/"

TEXT_REPLACEMENTS = {
    "love-and-imperfection": {
        "無理的行為讓在場女嘉賓感到傻眼。": "無禮的行為讓在場女嘉賓感到傻眼。",
        "他會需要尋找各式不同程度別人對自己的肯定": "他會需要尋找來自不同對象、不同程度的肯定",
    },
    "boundary-06": {
        "許多傷害發生的，不是越界的當下，而是那段你還以為「只是關心」的期間。":
            "許多傷害並非只發生在越界的當下，也發生在那段你還以為「只是關心」的期間。",
    },
    "workplace-01": {
        "下周的報告還沒有做完": "下週的報告還沒有做完",
        "可以完美的兼顧工作和家庭": "可以完美地兼顧工作和家庭",
    },
    "post-election-self-care": {
        "壓力賀爾蒙": "壓力荷爾蒙",
        "延續1~2週都無法恢復": "延續 1～2 週都無法恢復",
    },
    "addiction-05": {
        "她甚至已經想好怎麼回復": "她甚至已經想好怎麼回覆",
    },
    "workplace-02": {
        "日復一日的對他說出這些話": "日復一日地對他說出這些話",
    },
    "boundary-05": {
        "不需要靠沈默或偽裝來保護自己": "不需要靠沉默或偽裝來保護自己",
    },
    "boundary-07": {
        "在這個的社會文化下": "在這樣的社會文化下",
        "tab換行可以用": "",
    },
    "workplace-08": {
        "以下有些建議：": "以下提供一些建議：",
    },
}

EXACT_PARAGRAPH_REPLACEMENTS = {
    "workplace-06": {
        "世界衛生組織（WHO）將職業倦怠（burnout）列入了國際疾病分類（International Classification of Diseases，簡稱ICD）中。這一決定旨在提高對職業倦怠的認識和重視，並促使政府採取更有效的措施來預防和處理職業倦怠。職業倦怠主要表現為三大症狀：":
            "世界衛生組織將職業倦怠列入 ICD-11，分類為與工作相關的「職業現象」，並非醫療疾病。職業倦怠主要表現為三大症狀：",
    },
    "addiction-07": {
        "研究指出，過度依賴 AI 協作會導致大腦與思考相關的神經活動下降，甚至無法清楚記得自己剛完成的內容。這就像是一塊正在發育的肌肉，如果每次用力前都有人幫你分擔重量，肌肉不會壞掉，但會慢慢萎縮。我們跳過了思考的陣痛期，也同時跳過了深層內化的機會。":
            "部分初步研究指出，過度依賴 AI 協作可能伴隨與思考相關的神經活動降低，也可能較難清楚記得自己剛完成的內容。這就像是一塊正在發育的肌肉，如果每次用力前都有人幫你分擔重量，肌肉不會壞掉，但會慢慢萎縮。我們跳過了思考的陣痛期，也同時跳過了深層內化的機會。",
    },
    "boundary-03": {
        "️保留對方訊息、通話紀錄、錄音或影像證據":
            "️保留對方訊息、通話紀錄及其他合法取得的證據；錄音或錄影是否適法會依情境而異，必要時請先詢問警察或法律專業人員",
        "️通報學校、公司或請求法律協助（如向警局報案申請保護令）":
            "️通報學校、公司或向警察報案；是否能聲請保護令及適用程序，請由警察、法院或法律專業人員依個案說明",
    },
    "workplace-05": {
        "盡可能蒐集證據，例如記錄霸凌內容、時間，如果被加諸不合理的工作負擔，也可記錄下每日的工作項目及細節。另外事發當下也可以善用手機等行動裝置錄音、錄影，記錄自己被惡意對待的過程，若未來採取法律行動時這些行動都將成為有力的證明。":
            "盡可能保留事件內容、時間、工作指派及相關書面紀錄。在不危及自身安全且符合法令的前提下，其他蒐證方式可先詢問警察、主管機關或法律專業人員；個別資料能否作為證據，仍須依實際程序判斷。",
    },
    "boundary-04": {
        "️衛生福利部性影像處理中心－私ME專線：02-66057373 (每日上午9時至晚上10時，全年無休）。":
            "️如需性影像或人身安全相關協助，可撥打政府 113 保護專線，或使用社會安全網「關懷 e 起來」線上求助",
        "️撥打113專線(24小時免付費求助諮詢電話)，或「社會安全網-關懷e起來」線上諮詢([連結])。":
            f"️撥打 113 保護專線，或使用社會安全網「關懷 e 起來」線上求助：{ECARE_URL}",
        "️婦女救援基金會-未得同意散布性私密影像求助諮詢專線：02-2555-8595": "",
        "️現代婦女基金會-性騷擾諮詢專線：02-2351-2811": "",
        "️勵馨基金會：北部 02-8911-5595 #122；中部 04-2223-9595；南部 07-2237-955": "",
        "️台灣展翅協會：02-2562123 分機 280": "",
        "️法扶全國專線：市話412-8518；手機02-4128518": "",
    },
    "relationship-control": {
        "️撥打24小時保護專線「113」進行諮詢與通報":
            "️撥打政府 113 保護專線進行諮詢或求助",
        "️現代婦女基金會提供被害人支持服務，專線：(02)7728-5098 分機 6": "",
        "️勵馨基金會提供被害人庇護安置等服務，專線：北部 02-8911-5595 #122；中部 04-2223-9595；南部 07-2237-955": "",
    },
}


def compact(text: str) -> str:
    return " ".join(text.replace("\u3000", " ").split())


def normalize_title(text: str) -> str:
    text = text.replace("沒官熙~真的沒關係~~", "沒官熙～真的沒關係～～")
    text = re.sub(
        r"(?<=[\u4e00-\u9fff」』])\s*[-—－]\s*(?=[\u4e00-\u9fff「『])",
        "——",
        text,
    )
    return add_latin_spacing(text)


def add_latin_spacing(text: str) -> str:
    text = re.sub(r"(?<=[\u4e00-\u9fff])(?=[A-Za-z])", " ", text)
    text = re.sub(r"(?<=[A-Za-z])(?=[\u4e00-\u9fff])", " ", text)
    return text


def apply_approved_edits(slug: str, text: str) -> str:
    if slug == "post-election-self-care" and (
        "小李滿懷期待A候選人" in text or "候選人竟然敗選了…，" in text
    ):
        return text
    if slug == "addiction-05" and "「看一下就好。」，滑完IG" in text:
        return text
    if slug == "addiction-07" and "容易因其「零挫折」的特性" in text:
        return text
    if slug == "boundary-04":
        if "私ME專線" in text:
            return "● 如需性影像或人身安全相關協助，可撥打政府 113 保護專線，或使用社會安全網「關懷 e 起來」線上求助"
        if "[連結]" in text:
            return f"● 撥打 113 保護專線，或使用社會安全網「關懷 e 起來」線上求助：{ECARE_URL}"
        if any(name in text for name in ("婦女救援基金會-", "現代婦女基金會-", "勵馨基金會：", "台灣展翅協會：", "法扶全國專線：")):
            return ""
    if slug == "relationship-control":
        if "撥打24小時保護專線" in text:
            return "● 撥打政府 113 保護專線進行諮詢或求助"
        if "現代婦女基金會提供" in text or "勵馨基金會提供" in text:
            return ""
    if slug == "boundary-03":
        if "保留對方訊息、通話紀錄、錄音或影像證據" in text:
            return "● 保留對方訊息、通話紀錄及其他合法取得的證據；錄音或錄影是否適法會依情境而異，必要時請先詢問警察或法律專業人員"
        if "向警局報案申請保護令" in text:
            return "● 通報學校、公司或向警察報案；是否能聲請保護令及適用程序，請由警察、法院或法律專業人員依個案說明"
    if slug == "workplace-05" and text.startswith("盡可能蒐集證據"):
        return "盡可能保留事件內容、時間、工作指派及相關書面紀錄。在不危及自身安全且符合法令的前提下，其他蒐證方式可先詢問警察、主管機關或法律專業人員；個別資料能否作為證據，仍須依實際程序判斷。"
    if text in EXACT_PARAGRAPH_REPLACEMENTS.get(slug, {}):
        text = EXACT_PARAGRAPH_REPLACEMENTS[slug][text]
    for original, replacement in TEXT_REPLACEMENTS.get(slug, {}).items():
        text = text.replace(original, replacement)
    text = add_latin_spacing(text)
    text = re.sub(r"(?<=[\u4e00-\u9fff])(\d+(?:～\d+)?)(?=[\u4e00-\u9fff])", r" \1 ", text)
    if not any(mark in text for mark in ("「", "」", "『", "』", "“", "”")):
        text = re.sub(r"(?:\.{3,}|…+)", "……", text)
    return re.sub(r"\s{2,}", " ", text).strip()


def escape_with_links(text: str) -> str:
    escaped = html.escape(text)
    return re.sub(
        r"https://ecare\.mohw\.gov\.tw/",
        '<a href="https://ecare.mohw.gov.tw/" target="_blank" rel="noopener noreferrer">https://ecare.mohw.gov.tw/</a>',
        escaped,
    )


def clean_generated_html(text: str) -> str:
    return "\n".join(line.rstrip() for line in text.splitlines()) + "\n"


def series_image(series: Series, extension: str = "png") -> str:
    return f"{SITE}/assets/illustrations/{series.image_stem}-1600x900.{extension}"


def is_separator(text: str) -> bool:
    return bool(re.fullmatch(r"[-─—_]{8,}", text))


def paragraph_size(paragraph) -> float:
    sizes = [run.font.size.pt for run in paragraph.runs if run.font.size]
    return max(sizes, default=0)


def classify_title(source_title: str, other_count: int) -> tuple[str, int | None, str]:
    patterns = [
        ("addiction", r"「癮」未條\s*[#＃]\s*(\d+)\s*[-—－:]?\s*(.*)"),
        ("workplace", r"咖啡哪有工作苦\s*[#＃]\s*(\d+)\s*[-—－:]?\s*(.*)"),
        ("boundary", r"「性騷擾」是什麼\?\s*[#＃]\s*(\d+)\s*[-—－:]?\s*(.*)"),
        ("self-care", r"心理師的心靈雞湯\s*[#＃]\s*(\d+)\s*[-—－:]?\s*(.*)"),
    ]
    normalized = source_title.replace("？", "?")
    for key, pattern in patterns:
        match = re.match(pattern, normalized, flags=re.I)
        if match:
            number = int(match.group(1))
            display = match.group(2).strip(" -—－:")
            return key, number, display
    other_slugs = [
        "love-and-imperfection",
        "post-election-self-care",
        "relationship-control",
    ]
    return "insights", None, source_title, other_slugs[other_count]


def is_heading(paragraph, text: str, next_text: str) -> bool:
    if len(text) > 56 or text.endswith(("。", "！", "？", "；", "，", ".", "!", "?", "：", ":")):
        return False
    if text.startswith(("「", "『", "“", "\"", "●", "•", "▪", "-", "—")):
        return False
    bold_chars = sum(len(run.text) for run in paragraph.runs if run.bold)
    if bold_chars >= max(2, len(text) // 2):
        return True
    return bool(next_text and len(next_text) >= len(text) * 1.8 and len(text) <= 28)


def render_body(paragraphs, slug: str) -> tuple[str, str]:
    cleaned = []
    for paragraph in paragraphs:
        text = apply_approved_edits(slug, compact(paragraph.text))
        if re.match(
            r"(?:「癮」未條|咖啡哪有工作苦|「性騷擾」是什麼[?？]?|心理師的心靈雞湯)\s*[#＃]\s*\d+",
            text,
        ):
            continue
        if text and "黃郁倩" not in text and not is_separator(text):
            cleaned.append((paragraph, text))

    rendered = []
    list_items = []
    description = ""

    def flush_list():
        if list_items:
            rendered.append("<ul>\n" + "\n".join(list_items) + "\n</ul>")
            list_items.clear()

    for index, (paragraph, text) in enumerate(cleaned):
        next_text = cleaned[index + 1][1] if index + 1 < len(cleaned) else ""
        escaped = escape_with_links(text)
        if text.startswith(("●", "•", "▪", "▫")):
            list_items.append(f"  <li>{escape_with_links(text.lstrip('●•▪▫ ').strip())}</li>")
            continue
        flush_list()
        if is_heading(paragraph, text, next_text):
            rendered.append(f"<h2>{escaped}</h2>")
        elif text.startswith(("「", "『", "“", "\"")) and len(text) <= 90:
            rendered.append(f"<blockquote>{escaped}</blockquote>")
        else:
            rendered.append(f"<p>{escaped}</p>")
            if not description and len(text) >= 35:
                description = text[:115].rstrip("，、；： ") + ("…" if len(text) > 115 else "")
    flush_list()
    return "\n\n".join(rendered), description


def page_header(active: str = "articles") -> str:
    current = ' aria-current="page"' if active == "articles" else ""
    return f"""<header class="site-header">
    <div class="nav-wrap">
      <a class="brand" href="../index.html">{AUTHOR}<small>諮心字第005821號</small></a>
      <nav class="desktop-nav" aria-label="主選單">
        <a href="../index.html">首頁</a>
        <a href="../index.html#about">關於我</a>
        <a href="../index.html#services">諮商服務</a>
        <a href="../articles.html"{current}>心理專欄</a>
        <a href="../index.html#contact">聯絡預約</a>
      </nav>
      <a class="btn btn-primary nav-cta" href="{LINE_URL}" target="_blank" rel="noopener noreferrer">LINE 預約</a>
    </div>
  </header>"""


def footer() -> str:
    return f"""<footer class="site-footer">
    <div class="footer-inner">
      <div><strong>{AUTHOR}</strong><p>諮心字第005821號</p></div>
      <div class="footer-links">
        <a href="../articles.html">心理專欄</a>
        <a href="../index.html#services">諮商服務</a>
        <a href="../index.html#contact">聯絡預約</a>
      </div>
    </div>
  </footer>"""


def article_page(article, articles) -> str:
    series = SERIES[article["series"]]
    same_series = [item for item in articles if item["series"] == article["series"]]
    position = same_series.index(article)
    previous = same_series[position - 1] if position > 0 else None
    following = same_series[position + 1] if position + 1 < len(same_series) else None
    related = [item for item in same_series if item is not article][:3]
    canonical = f"{SITE}/articles/{article['slug']}.html"
    article_json_ld = {
        "@context": "https://schema.org",
        "@type": "Article",
        "headline": article["title"],
        "description": article["description"],
        "image": series_image(series),
        "url": canonical,
        "author": {"@type": "Person", "name": AUTHOR, "url": f"{SITE}/#about"},
        "publisher": {
            "@type": "Organization",
            "name": AUTHOR,
            "logo": {"@type": "ImageObject", "url": f"{SITE}/assets/headshot.jpg"},
        },
        "mainEntityOfPage": {"@type": "WebPage", "@id": canonical},
    }
    breadcrumb_json_ld = {
        "@context": "https://schema.org",
        "@type": "BreadcrumbList",
        "itemListElement": [
            {
                "@type": "ListItem",
                "position": 1,
                "name": "首頁",
                "item": f"{SITE}/",
            },
            {
                "@type": "ListItem",
                "position": 2,
                "name": "心理專欄",
                "item": f"{SITE}/articles.html",
            },
            {
                "@type": "ListItem",
                "position": 3,
                "name": series.name,
                "item": f"{SITE}/series/{series.key}.html",
            },
            {
                "@type": "ListItem",
                "position": 4,
                "name": article["title"],
                "item": canonical,
            },
        ],
    }
    series_label = series.name + (f" #{article['number']}" if article["number"] else "")

    def nav_item(label, item):
        if not item:
            return f'<div class="post-nav-item disabled"><span>{label}</span><strong>沒有其他文章</strong></div>'
        return (
            f'<a class="post-nav-item" href="{item["slug"]}.html"><span>{label}</span>'
            f'<strong>{html.escape(item["title"])}</strong></a>'
        )

    related_html = "\n".join(
        f'<a href="{item["slug"]}.html"><span>{html.escape(series.name)}</span>'
        f'<strong>{html.escape(item["title"])}</strong></a>'
        for item in related
    )
    return f"""<!doctype html>
<html lang="zh-Hant">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{html.escape(article["title"])}｜{html.escape(series.name)}｜{AUTHOR}</title>
  <meta name="description" content="{html.escape(article["description"], quote=True)}">
  <meta name="author" content="{AUTHOR}">
  <meta name="robots" content="index, follow">
  <link rel="canonical" href="{canonical}">
  <meta property="og:type" content="article">
  <meta property="og:title" content="{html.escape(article["title"], quote=True)}">
  <meta property="og:description" content="{html.escape(article["description"], quote=True)}">
  <meta property="og:image" content="{series_image(series)}">
  <meta property="og:image:width" content="1600">
  <meta property="og:image:height" content="900">
  <meta property="og:image:alt" content="{html.escape(series.image_alt, quote=True)}">
  <meta property="og:url" content="{canonical}">
  <meta property="og:locale" content="zh_TW">
  <meta name="twitter:card" content="summary_large_image">
  <meta name="twitter:title" content="{html.escape(article["title"], quote=True)}">
  <meta name="twitter:description" content="{html.escape(article["description"], quote=True)}">
  <meta name="twitter:image" content="{series_image(series)}">
  <meta name="twitter:image:alt" content="{html.escape(series.image_alt, quote=True)}">
  <link rel="stylesheet" href="../article.css?v=20260613-2">
  <script type="application/ld+json">{json.dumps(article_json_ld, ensure_ascii=False, indent=2)}</script>
  <script type="application/ld+json">{json.dumps(breadcrumb_json_ld, ensure_ascii=False, indent=2)}</script>
</head>
<body>
  {page_header()}
  <main>
    <article class="article-page">
      <header class="article-hero">
        <div class="page-inner narrow">
          <nav class="breadcrumb" aria-label="麵包屑">
            <a href="../index.html">首頁</a><span>/</span>
            <a href="../articles.html">心理專欄</a><span>/</span>
            <a href="../series/{series.key}.html">{html.escape(series.name)}</a>
          </nav>
          <p class="article-series">{html.escape(series_label)}</p>
          <h1>{html.escape(article["title"])}</h1>
          <p class="article-description">{html.escape(article["description"])}</p>
          <div class="article-meta">
            <span>作者：{AUTHOR}</span><span>分類：{html.escape(series.topic)}</span>
          </div>
        </div>
      </header>
      <div class="page-inner article-layout">
        <aside class="article-sidebar">
          <div class="author-card">
            <img src="../assets/headshot.jpg" alt="黃郁倩諮商心理師照片">
            <strong>{AUTHOR}</strong><span>諮心字第005821號</span>
            <p>陪伴兒童、青少年與成人面對情緒、關係、壓力與自我探索議題。</p>
            <a href="../index.html#about">了解心理師</a>
          </div>
          <div class="series-card">
            <span>目前分類</span><strong>{html.escape(series.name)}</strong>
            {f'<small>{html.escape(series.subtitle)}</small>' if series.subtitle else ''}
            <p>{html.escape(series.description)}</p>
            <a href="../series/{series.key}.html">回到分類頁</a>
          </div>
        </aside>
        <div class="article-content">
          {article["body"]}
          <aside class="article-disclaimer" aria-label="內容聲明">
            <strong>內容聲明</strong>
            <p>{DISCLAIMER}</p>
          </aside>
          <section class="article-cta">
            <p class="section-kicker">想進一步談談嗎？</p>
            <h2>閱讀可以是一個開始，諮商能陪你更靠近自己的需要。</h2>
            <p>如果文章裡的情境讓你覺得熟悉，可以為自己保留一個被理解與整理的空間。</p>
            <a class="btn btn-primary" href="{LINE_URL}" target="_blank" rel="noopener noreferrer">LINE 預約諮商</a>
          </section>
          <nav class="post-nav" aria-label="上一篇與下一篇">
            {nav_item("上一篇", previous)}
            {nav_item("下一篇", following)}
          </nav>
          <section class="related-section">
            <div class="section-head compact"><p class="section-kicker">延伸閱讀</p><h2>同分類文章</h2></div>
            <div class="related-grid">{related_html}</div>
          </section>
        </div>
      </div>
    </article>
  </main>
  {footer()}
</body>
</html>
"""


def series_page(series: Series, items, all_series) -> str:
    cards = "\n".join(
        f"""<article class="episode-card">
          <span>{html.escape(series.name)}{f" #{item['number']}" if item['number'] else ""}</span>
          <h3>{html.escape(item["title"])}</h3>
          <p>{html.escape(item["description"])}</p>
          <a href="../articles/{item["slug"]}.html">閱讀文章</a>
        </article>"""
        for item in items
    )
    related = "\n".join(
        f"""<a href="{other.key}.html"><span>{html.escape(other.topic)}</span>
          <strong>{html.escape(other.name)}</strong><p>{html.escape(other.description)}</p></a>"""
        for other in all_series
        if other.key != series.key
    )[:]
    canonical = f"{SITE}/series/{series.key}.html"
    collection_json_ld = {
        "@context": "https://schema.org",
        "@type": "CollectionPage",
        "name": f"{series.name}｜{series.topic}",
        "description": series.description,
        "url": canonical,
        "image": series_image(series),
        "author": {
            "@type": "Person",
            "name": AUTHOR,
            "url": f"{SITE}/#about",
        },
        "hasPart": [
            {
                "@type": "Article",
                "headline": item["title"],
                "url": f"{SITE}/articles/{item['slug']}.html",
            }
            for item in items
        ],
    }
    breadcrumb_json_ld = {
        "@context": "https://schema.org",
        "@type": "BreadcrumbList",
        "itemListElement": [
            {
                "@type": "ListItem",
                "position": 1,
                "name": "首頁",
                "item": f"{SITE}/",
            },
            {
                "@type": "ListItem",
                "position": 2,
                "name": "心理專欄",
                "item": f"{SITE}/articles.html",
            },
            {
                "@type": "ListItem",
                "position": 3,
                "name": series.name,
                "item": canonical,
            },
        ],
    }
    return f"""<!doctype html>
<html lang="zh-Hant">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{html.escape(series.name)}｜{html.escape(series.topic)}｜{AUTHOR}</title>
  <meta name="description" content="{html.escape(series.description, quote=True)}">
  <meta name="author" content="{AUTHOR}">
  <meta name="robots" content="index, follow">
  <link rel="canonical" href="{canonical}">
  <meta property="og:type" content="website">
  <meta property="og:title" content="{html.escape(series.name, quote=True)}｜{html.escape(series.topic, quote=True)}">
  <meta property="og:description" content="{html.escape(series.description, quote=True)}">
  <meta property="og:image" content="{series_image(series)}">
  <meta property="og:image:width" content="1600">
  <meta property="og:image:height" content="900">
  <meta property="og:image:alt" content="{html.escape(series.image_alt, quote=True)}">
  <meta property="og:url" content="{canonical}">
  <meta property="og:locale" content="zh_TW">
  <meta name="twitter:card" content="summary_large_image">
  <meta name="twitter:title" content="{html.escape(series.name, quote=True)}｜{html.escape(series.topic, quote=True)}">
  <meta name="twitter:description" content="{html.escape(series.description, quote=True)}">
  <meta name="twitter:image" content="{series_image(series)}">
  <meta name="twitter:image:alt" content="{html.escape(series.image_alt, quote=True)}">
  <link rel="stylesheet" href="../series.css?v=20260613-2">
  <script type="application/ld+json">{json.dumps(collection_json_ld, ensure_ascii=False, indent=2)}</script>
  <script type="application/ld+json">{json.dumps(breadcrumb_json_ld, ensure_ascii=False, indent=2)}</script>
</head>
<body>
  {page_header()}
  <main>
    <section class="series-hero">
      <div class="page-inner hero-grid">
        <div>
          <p class="eyebrow">{html.escape(series.topic)}</p>
          <h1>{html.escape(series.name)}</h1>
          {f'<p class="series-subtitle">{html.escape(series.subtitle)}</p>' if series.subtitle else ''}
          <p class="hero-subtitle">{html.escape(series.description)}</p>
          <div class="hero-actions">
            <a class="btn btn-primary" href="#all-episodes">查看全部文章</a>
            <a class="btn btn-secondary" href="../articles.html">回到心理專欄</a>
          </div>
        </div>
        <div class="series-visual">
          <picture class="series-hero-art">
            <source media="(max-width: 639px)" srcset="../assets/illustrations/{series.image_stem}-mobile-900x900.webp">
            <img src="../assets/illustrations/{series.image_stem}-1600x900.webp"
              alt="{html.escape(series.image_alt, quote=True)}" width="1600" height="900"
              fetchpriority="high" decoding="async">
          </picture>
          <aside class="series-note"><span>閱讀提示</span><p>{html.escape(series.note)}</p></aside>
        </div>
      </div>
    </section>
    <section class="section" id="all-episodes">
      <div class="page-inner">
        <div class="section-head">
          <p class="section-kicker">文章列表</p>
          <h2>共 {len(items)} 篇文章</h2>
          <p>以下文章皆可獨立閱讀；有編號的系列可依序閱讀。</p>
        </div>
        <div class="episode-list">{cards}</div>
      </div>
    </section>
    <section class="section alt">
      <div class="page-inner">
        <div class="section-head"><p class="section-kicker">其他分類</p><h2>繼續探索心理專欄</h2></div>
        <div class="related-grid">{related}</div>
      </div>
    </section>
    <section class="section">
      <div class="page-inner author-panel">
        <img src="../assets/headshot.jpg" alt="黃郁倩諮商心理師照片">
        <div><p class="section-kicker">作者</p><h2>{AUTHOR}</h2>
          <p>諮心字第005821號。提供兒童、青少年與成人心理諮商。</p>
          <div class="author-links">
            <a class="btn btn-secondary" href="../index.html#about">了解心理師</a>
            <a class="btn btn-primary" href="{LINE_URL}" target="_blank" rel="noopener noreferrer">LINE 預約諮商</a>
          </div>
        </div>
      </div>
    </section>
  </main>
  {footer()}
</body>
</html>
"""


def inventory_markdown(articles) -> str:
    rows = [
        "# 專欄文章盤點表",
        "",
        f"- 來源：`{SOURCE}`",
        f"- 總篇數：{len(articles)}",
        "- 作者顯示：黃郁倩 諮商心理師",
        "- 注意：標題與正文疑似錯字未於工程階段修改。",
        "",
        "| # | 分類 | 原始標題 | 網址 | 原文段落 |",
        "|---:|---|---|---|---|",
    ]
    for index, article in enumerate(articles, 1):
        rows.append(
            f"| {index} | {SERIES[article['series']].name} | "
            f"{article['source_title'].replace('|', '｜')} | "
            f"`/articles/{article['slug']}.html` | "
            f"{article['start'] + 1}–{article['end']} |"
        )
    return "\n".join(rows) + "\n"


def main():
    document = Document(SOURCE)
    title_indexes = [
        index
        for index, paragraph in enumerate(document.paragraphs)
        if compact(paragraph.text) and paragraph_size(paragraph) >= 23
    ]
    articles = []
    other_count = 0
    for position, start in enumerate(title_indexes):
        end = title_indexes[position + 1] if position + 1 < len(title_indexes) else len(document.paragraphs)
        source_title = compact(document.paragraphs[start].text)
        classified = classify_title(source_title, other_count)
        if len(classified) == 4:
            series_key, number, title, slug = classified
            other_count += 1
        else:
            series_key, number, title = classified
            slug = f"{series_key}-{number:02d}"
        body_paragraphs = list(document.paragraphs[start + 1 : end])
        while body_paragraphs:
            first_text = compact(body_paragraphs[0].text)
            duplicate_series_title = (
                number is not None
                and re.match(
                    rf"(?:心理師的心靈雞湯|{re.escape(SERIES[series_key].name)})\s*[#＃]\s*{number}",
                    first_text,
                )
            )
            if first_text in {source_title, title} or duplicate_series_title:
                body_paragraphs.pop(0)
                continue
            break
        title = normalize_title(title)
        body, description = render_body(body_paragraphs, slug)
        articles.append(
            {
                "series": series_key,
                "number": number,
                "title": title,
                "source_title": source_title,
                "slug": slug,
                "body": body,
                "description": description or title,
                "start": start,
                "end": end,
            }
        )

    articles.sort(key=lambda item: (list(SERIES).index(item["series"]), item["number"] or 99, item["start"]))
    article_dir = ROOT / "articles"
    series_dir = ROOT / "series"
    article_dir.mkdir(exist_ok=True)
    series_dir.mkdir(exist_ok=True)
    for article in articles:
        (article_dir / f"{article['slug']}.html").write_text(
            clean_generated_html(article_page(article, articles)),
            encoding="utf-8",
        )
    for series in SERIES.values():
        items = [article for article in articles if article["series"] == series.key]
        (series_dir / f"{series.key}.html").write_text(
            clean_generated_html(series_page(series, items, list(SERIES.values()))),
            encoding="utf-8",
        )
    (WORKFLOW / "project" / "article-inventory.md").write_text(
        inventory_markdown(articles),
        encoding="utf-8",
    )
    print(f"Generated {len(articles)} articles and {len(SERIES)} series pages.")


if __name__ == "__main__":
    main()
